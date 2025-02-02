from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('Document Title', 0)  #插入标题

p = document.add_paragraph('A plain paragraph having some ')   #插入段落
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='IntenseQuote')

document.add_paragraph(
    'first item in unordered list', style='ListBullet'
)
document.add_paragraph(
    'first item in ordered list', style='ListNumber'
)

document.add_picture(r'./common/IMG_3817.jpg', width=Inches(1.25)) #插入图片

recordset = [{'Qty': 10, 'Id': 1, 'Desc': 'Sugar'}, 
    {'Qty': 5, 'Id': 2, 'Desc': 'Salt'}]

table = document.add_table(rows=1, cols=3) #插入表格
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for item in recordset:
    row_cells = table.add_row().cells
    row_cells[0].text = str(item['Qty'])
    row_cells[1].text = str(item['Id'])
    row_cells[2].text = item['Desc']

document.add_page_break()

document.save('demo.docx')  #保存文档
